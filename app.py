import streamlit as st
import pandas as pd
from datetime import datetime, timedelta
import gspread
from google.oauth2.service_account import Credentials
import os
from io import BytesIO
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import calendar
import time
import re
import itertools
import math
import altair as alt
from functools import wraps 
import extra_streamlit_components as stx 

# --- 1. CONFIGURACIÓN DE PÁGINA ---
st.set_page_config(
    page_title="El Estudio", 
    page_icon="🔥", 
    layout="wide", 
    initial_sidebar_state="collapsed" 
)

# --- INICIALIZAR EL SELLO VIP (COOKIES) ---
gestor_cookies = stx.CookieManager()

def cerrar_sesion_seguro():
    gestor_cookies.delete("sello_vip_estudio") 
    st.session_state['logueado'] = False 
    st.session_state['esperando_login_manual'] = True 
    st.query_params.clear() 
    st.rerun() 

# --- 2. ESTILOS CSS ---
def cargar_estilos():
    st.markdown("""
        <style>
        @import url('https://fonts.googleapis.com/css2?family=Montserrat:wght@300;400;600;800&display=swap');
        html, body, [class*="css"] { 
            font-family: 'Montserrat', sans-serif; 
            -webkit-text-size-adjust: 100%;
        }
        
        .block-container { padding-top: 1.5rem; padding-bottom: 5rem; }

        h1 { color: #E63946 !important; font-weight: 800 !important; letter-spacing: -1px; margin-bottom: 0.5rem; }
        h2, h3 { font-weight: 600 !important; margin-top: 1rem; }
        
        /* Botones */
        div.stButton > button:first-child {
            background-color: #E63946; color: white; border-radius: 12px; border: none;
            font-weight: 600; text-transform: uppercase; letter-spacing: 1px;
            padding: 18px 20px; width: 100%;
            transition: all 0.2s ease;
            box-shadow: 0 4px 6px rgba(0,0,0,0.1);
        }
        div.stButton > button:first-child:active { transform: scale(0.98); }
        
        /* Inputs */
        input, textarea, select, div[data-baseweb="select"] { 
            font-size: 16px !important; 
            border-radius: 8px !important; 
            min-height: 45px;
        }
        
        /* Estilo de Expander (Tarjetas) */
        .streamlit-expanderHeader {
            font-weight: 600;
            background-color: #262730;
            border-radius: 8px;
        }
        
        /* Calendario */
        .calendar-container {
            display: grid; grid-template-columns: repeat(7, 1fr); gap: 4px; margin-top: 10px; margin-bottom: 20px;
        }
        .calendar-day {
            aspect-ratio: 1; display: flex; align-items: center; justify-content: center;
            border-radius: 6px; font-weight: bold; font-size: 0.85rem; color: #FFF; background-color: #262730;
        }
        .day-completed { background-color: #2ECC71 !important; color: #000 !important; }
        .day-incomplete { background-color: #F39C12 !important; color: #000 !important; }
        
        div[data-testid="stVerticalBlock"] { gap: 1rem; }

        #MainMenu {visibility: hidden !important;}
        footer {visibility: hidden !important; display: none !important;}
        header {visibility: hidden !important;}
        
        [data-testid="stToolbar"] { visibility: hidden !important; display: none !important; }
        [data-testid="stHeader"] { visibility: hidden !important; background: transparent !important; }
        .stAppDeployButton { display: none !important; visibility: hidden !important; }

        .big-metric { font-size: 3rem; font-weight: 800; color: #E63946; text-align: center; }
        .sub-metric { font-size: 1rem; color: #aaa; text-align: center; margin-bottom: 20px; }

        div[data-testid="stVerticalBlockBorderWrapper"]:has(.red-border-trigger) {
            border: 3px solid #E63946 !important; border-radius: 10px !important;
        }
        
        [data-testid="stArrowVegaLiteChart"], [data-testid="stVegaLiteChart"], canvas.marks {
            pointer-events: none !important;
        }
        
        /* Ajuste de tablas en markdown */
        table { width: 100%; border-collapse: collapse; }
        th, td { padding: 10px !important; text-align: center !important; }
        </style>
    """, unsafe_allow_html=True)

cargar_estilos()

# --- ARMADURA DE REINTENTOS ---
def reintentar_conexion(intentos=3, espera=2):
    def decorador(funcion):
        @wraps(funcion)
        def envoltura(*args, **kwargs):
            for intento in range(intentos):
                try:
                    return funcion(*args, **kwargs)
                except Exception as e:
                    if intento < intentos - 1:
                        time.sleep(espera)
                    else:
                        st.error("⚠️ Tu conexión a internet parece inestable. Por favor, intenta de nuevo.")
                        raise e 
        return envoltura
    return decorador

# --- 3. CONEXIÓN GOOGLE SHEETS ---
SCOPES = ["https://www.googleapis.com/auth/spreadsheets", "https://www.googleapis.com/auth/drive"]

@st.cache_resource
@reintentar_conexion()
def conectar_google_sheet():
    if os.path.exists("mis_secretos.json"):
        creds = Credentials.from_service_account_file("mis_secretos.json", scopes=SCOPES)
    else:
        creds_dict = st.secrets["gcp_service_account"]
        creds = Credentials.from_service_account_info(creds_dict, scopes=SCOPES)
    return gspread.authorize(creds).open("El Estudio DB")

def natural_sort_key(s):
    return [int(text) if text.isdigit() else text.lower() for text in re.split('([0-9]+)', str(s))]

# --- 4. FUNCIONES DE LECTURA ---
@st.cache_data(ttl=600)
@reintentar_conexion() 
def obtener_todos_usuarios():
    sh = conectar_google_sheet()
    df = pd.DataFrame(sh.worksheet("Usuarios").get_all_records())
    if df.empty: return df
    df.columns = [str(c).strip().capitalize() for c in df.columns]
    df = df.loc[:, ~df.columns.duplicated()]
    if "" in df.columns: df = df.drop(columns=[""])
    
    if "Rol" in df.columns: df["Rol"] = df["Rol"].astype(str).str.strip().str.lower()
    if "Usuario" in df.columns: df["Usuario"] = df["Usuario"].astype(str).str.strip()
    return df

@st.cache_data(ttl=600)
@reintentar_conexion() 
def leer_rutina(alumno):
    sh = conectar_google_sheet()
    df = pd.DataFrame(sh.worksheet("Rutinas").get_all_records())
    if df.empty: return df
    df.columns = [str(c).strip().capitalize() for c in df.columns]
    df = df.loc[:, ~df.columns.duplicated()]
    if "" in df.columns: df = df.drop(columns=[""])
    
    df["Alumno_Norm"] = df["Alumno"].astype(str).str.strip().str.lower()
    alumno_norm = alumno.strip().lower()
    return df[df["Alumno_Norm"] == alumno_norm].drop(columns=["Alumno_Norm"])

@st.cache_data(ttl=600)
@reintentar_conexion()
def leer_sesiones_alumno(alumno):
    sh = conectar_google_sheet()
    df = pd.DataFrame(sh.worksheet("Sesiones").get_all_records())
    if df.empty: return df
    df.columns = [str(c).strip().capitalize() for c in df.columns]
    df = df.loc[:, ~df.columns.duplicated()]
    if "" in df.columns: df = df.drop(columns=[""])
    
    if "Usuario" not in df.columns: return pd.DataFrame()
    df["Usuario_Norm"] = df["Usuario"].astype(str).str.strip().str.lower()
    alumno_norm = alumno.strip().lower()
    df_alumno = df[df["Usuario_Norm"] == alumno_norm].copy()
    if not df_alumno.empty and "Fecha" in df_alumno.columns:
        df_alumno["Fecha"] = pd.to_datetime(df_alumno["Fecha"], format='mixed', errors='coerce').dt.normalize()
    return df_alumno

@st.cache_data(ttl=600)
@reintentar_conexion() 
def leer_registros_alumno(alumno):
    sh = conectar_google_sheet()
    df = pd.DataFrame(sh.worksheet("Registros").get_all_records())
    if df.empty: return df
    col_map = {c: str(c).strip().capitalize() for c in df.columns}
    for k, v in col_map.items():
        if v.lower() in ['user', 'alumno']: col_map[k] = 'Usuario'
        if v.lower() in ['date', 'day']: col_map[k] = 'Fecha'
    df = df.rename(columns=col_map)
    
    df = df.loc[:, ~df.columns.duplicated()]
    if "" in df.columns: df = df.drop(columns=[""])
    
    if "Usuario" not in df.columns: return pd.DataFrame()
    df["Usuario_Norm"] = df["Usuario"].astype(str).str.strip().str.lower()
    alumno_norm = alumno.strip().lower()
    df_alumno = df[df["Usuario_Norm"] == alumno_norm].copy()
    if not df_alumno.empty and "Fecha" in df_alumno.columns:
        df_alumno["Fecha"] = pd.to_datetime(df_alumno["Fecha"], format='mixed', errors='coerce').dt.normalize()
        df_alumno = df_alumno.dropna(subset=["Fecha"])
        df_alumno = df_alumno.reset_index(drop=True)
        def extraer_numero(texto):
            try: return float(re.search(r"(\d+[.,]?\d*)", str(texto)).group(1).replace(",", "."))
            except: return 0.0
        if "Peso" in df_alumno.columns: df_alumno["Peso_Grafico"] = df_alumno["Peso"].apply(extraer_numero)
        col_reps = next((c for c in df_alumno.columns if "rep" in c.lower()), None)
        if col_reps: df_alumno["Reps_Grafico"] = df_alumno[col_reps].apply(extraer_numero)
        else: df_alumno["Reps_Grafico"] = 0.0
    return df_alumno

def calcular_1rm_promedio(peso, reps):
    if reps == 1: return peso
    if reps == 0: return 0
    brzycki = peso / (1.0278 - 0.0278 * reps)
    epley = peso * (1 + 0.0333 * reps)
    lander = (100 * peso) / (101.3 - 2.67123 * reps)
    lombardi = peso * (reps ** 0.10)
    mayhew = (100 * peso) / (52.2 + (41.9 * math.exp(-0.055 * reps)))
    oconner = peso * (1 + 0.025 * reps)
    wathen = (100 * peso) / (48.8 + (53.8 * math.exp(-0.075 * reps)))
    promedio = (brzycki + epley + lander + lombardi + mayhew + oconner + wathen) / 7
    return promedio

# --- 5. FUNCIONES DE ESCRITURA (BLINDADAS) ---
def obtener_usuario(usuario_input, password_input):
    try:
        df = obtener_todos_usuarios()
        if df.empty: return None
        u_input = usuario_input.strip().lower()
        if "Usuario" not in df.columns: return None
        df["User_Lower"] = df["Usuario"].astype(str).str.strip().str.lower()
        usuario = df[(df["User_Lower"] == u_input) & (df["Password"].astype(str).str.strip() == password_input.strip())]
        return usuario.iloc[0] if not usuario.empty else None
    except: return None

def recuperar_usuario_por_nombre(usuario_input):
    try:
        df = obtener_todos_usuarios()
        if df.empty: return None
        if "Usuario" not in df.columns: return None
        u_input = usuario_input.strip().lower()
        df["User_Lower"] = df["Usuario"].astype(str).str.strip().str.lower()
        usuario = df[df["User_Lower"] == u_input]
        return usuario.iloc[0] if not usuario.empty else None
    except: return None

@reintentar_conexion() 
def guardar_rutina_actualizada(alumno, dia, df_c, df_f, df_ca):
    sh = conectar_google_sheet()
    ws = sh.worksheet("Rutinas")
    all_data = ws.get_all_records()
    df_old = pd.DataFrame(all_data)
    
    if df_old.empty:
        raw_vals = ws.get_all_values()
        if len(raw_vals) > 1:
            st.error("🚨 SISTEMA DE SEGURIDAD: Error de lectura detectado en Google Sheets. Guardado bloqueado para proteger tus rutinas.")
            return

    cols = ["Alumno", "Dia", "Seccion", "Orden", "Ejercicio", "Link", "Series", "Reps", "Kg", "Notas"]
    nuevas_filas = []
    
    def procesar_df(df, seccion):
        df = df.fillna("")
        for _, row in df.iterrows():
            if str(row.get("Ejercicio","")).strip() or str(row.get("Notas","")).strip():
                nuevas_filas.append({
                    "Alumno": alumno, "Dia": dia, "Seccion": seccion,
                    "Orden": str(row.get("Orden","-")), "Ejercicio": str(row.get("Ejercicio","")),
                    "Link": str(row.get("Link","")), "Series": str(row.get("Series","")),
                    "Reps": str(row.get("Reps","")), "Kg": str(row.get("Kg","-")),
                    "Notas": str(row.get("Notas",""))
                })
    
    procesar_df(df_c, "Calentamiento")
    procesar_df(df_f, "Fuerza")
    procesar_df(df_ca, "Cardio")

    if not df_old.empty:
        df_old.columns = [str(c).strip().capitalize() for c in df_old.columns]
        df_old = df_old.loc[:, ~df_old.columns.duplicated()]
        if "" in df_old.columns: df_old = df_old.drop(columns=[""])
        mask = ~((df_old["Alumno"].astype(str).str.lower() == alumno.lower()) & (df_old["Dia"].astype(str) == dia))
        df_clean = df_old[mask]
        df_final = pd.concat([df_clean, pd.DataFrame(nuevas_filas)], ignore_index=True)
        
        if len(df_old) >= 10 and len(df_final) < (len(df_old) * 0.5):
            st.error("🚨 SISTEMA DE SEGURIDAD CRÍTICO: Se intentó borrar una gran cantidad de rutinas por accidente. Guardado bloqueado automáticamente.")
            return
    else:
        df_final = pd.DataFrame(nuevas_filas)

    df_final = df_final.fillna("")
    for c in cols: 
        if c not in df_final.columns: df_final[c] = ""
        
    ws.clear()
    ws.update([cols] + df_final[cols].values.tolist())

def guardar_desde_tarjetas(alumno, dia, rut_hoy, session_state):
    rows_c = []; rows_f = []; rows_ca = []
    for idx, row in rut_hoy.iterrows():
        seccion = row["Seccion"]
        ej_limpio = re.sub(r'[^a-zA-Z0-9]', '', str(row['Ejercicio']))
        key_kg = f"kg_{alumno}_{dia}_{idx}_{ej_limpio}"
        key_notas = f"notas_{alumno}_{dia}_{idx}_{ej_limpio}"
        
        nuevo_kg = str(session_state.get(key_kg, row.get("Kg", "")))
        nueva_nota = str(session_state.get(key_notas, row.get("Notas", "")))
        
        new_row = row.copy()
        new_row["Kg"] = nuevo_kg
        new_row["Notas"] = nueva_nota
        if seccion == "Calentamiento": rows_c.append(new_row)
        elif seccion == "Fuerza": rows_f.append(new_row)
        elif seccion == "Cardio": rows_ca.append(new_row)
    
    df_c = pd.DataFrame(rows_c) if rows_c else pd.DataFrame()
    df_f = pd.DataFrame(rows_f) if rows_f else pd.DataFrame()
    df_ca = pd.DataFrame(rows_ca) if rows_ca else pd.DataFrame()
    guardar_rutina_actualizada(alumno, dia, df_c, df_f, df_ca)

@reintentar_conexion() 
def guardar_registro(usuario, ejercicio, peso, reps, rpe, notas, fecha_input=None):
    sh = conectar_google_sheet()
    fecha = fecha_input.strftime("%Y-%m-%d") if fecha_input else datetime.now().strftime("%Y-%m-%d") 
    
    peso_str = str(peso).replace(".", ",")
        
    try: reps_num = int(reps)
    except: reps_num = 0
    
    sh.worksheet("Registros").append_row(
        [fecha, usuario, ejercicio, peso_str, reps_num, int(rpe), str(notas)],
        value_input_option="USER_ENTERED"
    )

@reintentar_conexion() 
def actualizar_registros_masivo(usuario, df_editado):
    sh = conectar_google_sheet()
    ws = sh.worksheet("Registros")
    all_data = ws.get_all_records()
    df_all = pd.DataFrame(all_data)
    if df_all.empty: return
    df_all.columns = [str(c).strip().capitalize() for c in df_all.columns]
    
    df_all = df_all.loc[:, ~df_all.columns.duplicated()]
    if "" in df_all.columns: df_all = df_all.drop(columns=[""])
    
    col_map = {c: c for c in df_all.columns}
    for c in df_all.columns:
        if c.lower() in ['user', 'alumno']: col_map[c] = 'Usuario'
    df_all = df_all.rename(columns=col_map)
    df_otros = df_all[df_all["Usuario"].astype(str).str.lower() != usuario.lower()]
    df_nuevos = df_editado.copy()
    
    if "Fecha" in df_nuevos.columns:
        df_nuevos["Fecha"] = pd.to_datetime(df_nuevos["Fecha"], format='mixed', errors='ignore').dt.strftime("%Y-%m-%d")
        
    def parse_number_to_comma_string(val):
        if pd.isna(val) or str(val).strip() == "": return ""
        return str(val).replace(".", ",")

    def parse_int(val):
        if pd.isna(val) or str(val).strip() == "": return ""
        try: return int(float(str(val).replace(",", ".")))
        except: return val

    if "Peso" in df_nuevos.columns:
        df_nuevos["Peso"] = df_nuevos["Peso"].apply(parse_number_to_comma_string)
    if "Repeticiones" in df_nuevos.columns:
        df_nuevos["Repeticiones"] = df_nuevos["Repeticiones"].apply(parse_int)
        
    df_final = pd.concat([df_otros, df_nuevos], ignore_index=True).fillna("").sort_values("Fecha")
    ws.clear()
    
    ws.update([df_final.columns.values.tolist()] + df_final.values.tolist(), value_input_option="USER_ENTERED")

# CTO FIX: Modificamos el guardado de sesiones para incluir el "Dia_rutina" (la cuarta columna en Sheets)
@reintentar_conexion() 
def guardar_estado_sesion(usuario, estado, fecha_dt=None, dia_rutina=""):
    sh = conectar_google_sheet()
    ws = sh.worksheet("Sesiones")
    fecha_str = fecha_dt.strftime("%Y-%m-%d") if fecha_dt else datetime.now().strftime("%Y-%m-%d")
    ws.append_row([fecha_str, usuario, estado, dia_rutina])

def preparar_df_editor(df_input, columnas, filas_minimas=4):
    if df_input.empty: df_input = pd.DataFrame(columns=columnas)
    for c in ["Series", "Reps", "Kg"]:
        if c in df_input.columns: df_input[c] = df_input[c].astype(str)
    filas_actuales = len(df_input)
    if filas_actuales < filas_minimas:
        extras = pd.DataFrame([[""] * len(columnas)] * (filas_minimas - filas_actuales), columns=columnas)
        df_input = pd.concat([df_input, extras], ignore_index=True)
    return df_input

def generar_word(alumno, df_rutina):
    doc = Document()
    p_titulo = doc.add_heading(f'RUTINA: {alumno.upper()}', 0)
    p_titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Generado: {datetime.now().strftime('%d/%m/%Y')}")
    doc.add_paragraph("-" * 50)
    dias_ordenados = sorted(df_rutina["Dia"].unique(), key=natural_sort_key)
    for dia in dias_ordenados:
        doc.add_heading(dia, level=1)
        rutina_dia = df_rutina[df_rutina["Dia"] == dia]
        emoji_map = {"Calentamiento": "🔥", "Fuerza": "🏋️‍♂️", "Cardio": "🏃‍♂️"}
        for sec in ["Calentamiento", "Fuerza", "Cardio"]:
            df_sec = rutina_dia[rutina_dia["Seccion"] == sec]
            if not df_sec.empty:
                titulo_sec = f"{emoji_map.get(sec, '')} {sec}"
                doc.add_heading(titulo_sec, level=2)
                cols_count = 5 if sec == "Fuerza" else 4
                table = doc.add_table(rows=1, cols=cols_count)
                table.style = 'Table Grid'
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = 'EJERCICIO'
                hdr_cells[1].text = 'SERIES'
                hdr_cells[2].text = 'REPS'
                if sec == "Fuerza":
                    hdr_cells[3].text = 'KG'
                    hdr_cells[4].text = 'NOTAS'
                else:
                    hdr_cells[3].text = 'NOTAS'
                for _, row in df_sec.iterrows():
                    row_cells = table.add_row().cells
                    nom_ej = str(row['Ejercicio'])
                    if sec == "Fuerza":
                        orden = str(row.get('Orden', '')).strip()
                        if orden and orden != '-': nom_ej = f"{orden}. {nom_ej}"
                    row_cells[0].text = nom_ej
                    row_cells[1].text = str(row.get('Series', ''))
                    row_cells[2].text = str(row.get('Reps', ''))
                    if sec == "Fuerza":
                        row_cells[3].text = str(row.get('Kg', ''))
                        row_cells[4].text = str(row.get('Notas', ''))
                    else:
                        row_cells[3].text = str(row.get('Notas', ''))
                doc.add_paragraph("")
    b = BytesIO(); doc.save(b); b.seek(0); return b

def render_calendar(year, month, df_sesiones):
    nombres_meses = ["", "Enero", "Febrero", "Marzo", "Abril", "Mayo", "Junio", "Julio", "Agosto", "Septiembre", "Octubre", "Noviembre", "Diciembre"]
    month_name = nombres_meses[month]
    cal = calendar.monthcalendar(year, month)
    asistencia_map = {}
    if not df_sesiones.empty:
        mask = (df_sesiones["Fecha"].dt.year == year) & (df_sesiones["Fecha"].dt.month == month)
        df_mes = df_sesiones[mask]
        for _, row in df_mes.iterrows(): asistencia_map[row["Fecha"].day] = row["Estado"]
    html = f"""<div style="text-align:center; font-weight:bold; font-size:1.1rem; color:#E63946; margin-bottom:8px;">{month_name} {year}</div>
    <div class="calendar-container">
    <div style="text-align:center; color:#888;">L</div><div style="text-align:center; color:#888;">M</div><div style="text-align:center; color:#888;">M</div><div style="text-align:center; color:#888;">J</div><div style="text-align:center; color:#888;">V</div><div style="text-align:center; color:#888;">S</div><div style="text-align:center; color:#888;">D</div>"""
    for week in cal:
        for day in week:
            if day == 0: html += '<div></div>'
            else:
                css = "calendar-day"
                if asistencia_map.get(day) == "Completado": css += " day-completed"
                elif asistencia_map.get(day) == "Incompleto": css += " day-incomplete"
                html += f'<div class="{css}">{day}</div>'
    html += "</div>"
    st.markdown(html, unsafe_allow_html=True)

def renderizar_bloque_seccion(titulo, df_seccion, alumno, dia):
    if df_seccion.empty: return
    st.markdown(f"#### {titulo}")
    def obtener_bloque(orden):
        o = str(orden).strip().upper()
        if len(o) > 0 and o[0].isalpha() and o != "-":
            return o[0] 
        return "SIN_BLOQUE"
    datos_con_indice = []
    for idx, row in df_seccion.iterrows():
        datos_con_indice.append((idx, row, obtener_bloque(row.get("Orden", ""))))
    for bloque, grupo in itertools.groupby(datos_con_indice, key=lambda x: x[2]):
        items = list(grupo)
        if bloque != "SIN_BLOQUE":
            with st.container(border=True):
                st.markdown("<div class='red-border-trigger'></div>", unsafe_allow_html=True)
                for idx, row, _ in items:
                    render_tarjeta_individual(idx, row, alumno, dia)
        else:
            for idx, row, _ in items:
                render_tarjeta_individual(idx, row, alumno, dia)

def render_tarjeta_individual(idx, row, alumno, dia):
    ej_nombre = row['Ejercicio']
    series_reps = f"{row.get('Series','?')} x {row.get('Reps','?')}"
    orden_prefix = ""
    if "Orden" in row and str(row["Orden"]).strip() not in ["", "-"]:
        orden_prefix = f"**{row['Orden']}** | "
    titulo_card = f"{orden_prefix}{ej_nombre} ({series_reps})"
    with st.expander(titulo_card):
        link = str(row.get('Link', '')).strip()
        if link: st.link_button("📺 Ver Video", link, use_container_width=True)
        c_kg, c_notas = st.columns([1, 2])
        
        ej_limpio = re.sub(r'[^a-zA-Z0-9]', '', str(ej_nombre))
        k_kg = f"kg_{alumno}_{dia}_{idx}_{ej_limpio}"
        k_notas = f"notas_{alumno}_{dia}_{idx}_{ej_limpio}"
        
        if k_kg in st.session_state: val_kg = st.session_state[k_kg]
        else:
            val_kg = str(row.get('Kg', ''))
            if val_kg == "nan": val_kg = ""
        if k_notas in st.session_state: val_notas = st.session_state[k_notas]
        else:
            val_notas = str(row.get('Notas', ''))
            if val_notas == "nan": val_notas = ""
            
        with c_kg: st.text_input("Kg Realizados", value=val_kg, key=k_kg)
        with c_notas: st.text_area("Notas / Instrucciones", value=val_notas, key=k_notas, height=68)

# --- 6. GESTIÓN DE LOGIN ---
if 'logueado' not in st.session_state: st.session_state['logueado'] = False
if 'esperando_login_manual' not in st.session_state: st.session_state['esperando_login_manual'] = False

if not st.session_state['logueado']:
    if st.session_state.get('esperando_login_manual', False):
        usuario_a_validar = None
    else:
        usuario_url = st.query_params.get("u", None)
        usuario_cookie = gestor_cookies.get(cookie="sello_vip_estudio")
        usuario_a_validar = usuario_url if usuario_url else usuario_cookie

    if usuario_a_validar:
        u_data = recuperar_usuario_por_nombre(usuario_a_validar)
        if u_data is not None:
            st.session_state['logueado'] = True
            st.session_state['usuario_info'] = u_data
            if str(u_data.get('Rol','')).lower() == 'alumno':
                st.query_params["u"] = u_data['Usuario']
        else: 
            st.query_params.clear()

if not st.session_state['logueado']:
    col1, col2, col3 = st.columns([1, 2, 1])
    with col2:
        st.markdown("<br><br><h1 style='text-align: center;'>EL ESTUDIO 🔥</h1><br>", unsafe_allow_html=True)
        with st.container(border=True):
            with st.form("login"):
                u = st.text_input("USUARIO"); p = st.text_input("CONTRASEÑA", type="password")
                st.markdown("<br>", unsafe_allow_html=True)
                if st.form_submit_button("ENTRAR", use_container_width=True):
                    user = obtener_usuario(u, p)
                    if user is not None: 
                        st.session_state['logueado'] = True
                        st.session_state['usuario_info'] = user
                        st.session_state['esperando_login_manual'] = False
                        
                        fecha_expiracion = datetime.now() + timedelta(days=30)
                        gestor_cookies.set("sello_vip_estudio", user['Usuario'], expires_at=fecha_expiracion)
                        
                        if str(user.get('Rol','')).lower() == 'alumno':
                            st.query_params["u"] = user['Usuario']
                        time.sleep(0.5) 
                        st.rerun()
                    else: st.error("❌ Error de credenciales")
else:
    datos = st.session_state['usuario_info']
    rol = str(datos.get('Rol', 'alumno')).strip()
    nombre = str(datos.get('Nombre', 'Usuario')).strip()
    alias = str(datos.get('Usuario', '')).strip()

    if rol == "admin":
        st.title("🎛️ PANEL DE CONTROL")
        tab1, tab2, tab3 = st.tabs(["📝 DISEÑO DE RUTINAS", "📊 ESTADÍSTICAS", "⚙️ PERFIL"])
        
        with tab1:
            st.markdown("### 👤 Seleccionar Alumno y Día")
            with st.container(border=True):
                us = obtener_todos_usuarios()
                als = us[us["Rol"] == "alumno"]["Usuario"].tolist()
                
                alu = st.selectbox("ALUMNO", als)
                dia = st.radio("DÍA", ["Día 1", "Día 2", "Día 3", "Día 4", "Día 5"], horizontal=True)

            rut = leer_rutina(alu)
            cols_c = ["Ejercicio", "Link", "Series", "Reps", "Notas"]
            cols_f = ["Orden", "Ejercicio", "Link", "Series", "Reps", "Kg", "Notas"]
            cols_ca = ["Ejercicio", "Link", "Series", "Reps", "Notas"]
            d_cal = pd.DataFrame(columns=cols_c)
            d_fue = pd.DataFrame(columns=cols_f)
            d_car = pd.DataFrame(columns=cols_ca)

            if not rut.empty:
                r_dia = rut[rut["Dia"] == dia]
                if not r_dia.empty:
                    c = r_dia[r_dia["Seccion"] == "Calentamiento"]
                    if not c.empty: d_cal = c[cols_c]
                    f = r_dia[r_dia["Seccion"] == "Fuerza"]
                    if not f.empty: d_fue = f[cols_f]; d_fue["Kg"] = d_fue["Kg"].astype(str)
                    ca = r_dia[r_dia["Seccion"] == "Cardio"]
                    if not ca.empty: d_car = ca[cols_ca]

            d_cal = preparar_df_editor(d_cal, cols_c, filas_minimas=4)
            d_fue = preparar_df_editor(d_fue, cols_f, filas_minimas=8)
            d_car = preparar_df_editor(d_car, cols_ca, filas_minimas=3)
            
            cfg_comun = {
                "Ejercicio": st.column_config.TextColumn("Ejercicio", width="medium", required=True),
                "Link": st.column_config.LinkColumn("📺", width="small"),
                "Series": st.column_config.TextColumn("Ser.", width="small"),
                "Reps": st.column_config.TextColumn("Rep.", width="small"),
                "Notas": st.column_config.TextColumn("Notas", width="medium")
            }
            cfg_fuerza = cfg_comun.copy()
            cfg_fuerza["Orden"] = st.column_config.TextColumn("Ord.", width="small")
            cfg_fuerza["Kg"] = st.column_config.TextColumn("Kg", width="small")
            
            st.markdown("---")
            
            c_g_top, c_d_top = st.columns([1, 1])
            with c_g_top:
                guardado_arriba = st.button("💾 GUARDAR RUTINA RÁPIDO", type="primary", use_container_width=True, key="btn_guardar_top")
            with c_d_top:
                if not rut.empty: 
                    st.download_button("📥 DESCARGAR WORD", generar_word(alu, rut), f"{alu}.docx", use_container_width=True, key="btn_word_top")

            with st.expander("✏️ EDITOR DE RUTINA", expanded=True):
                st.info("💡 **Tip CTO:** Las columnas están compactadas para que entren mejor en tu celular. No olvides presionar Guardar.")
                
                st.caption("🔥 CALENTAMIENTO")
                ed_c = st.data_editor(d_cal, num_rows="dynamic", use_container_width=True, key=f"c_{alu}_{dia}", column_config=cfg_comun, hide_index=True)
                
                st.caption("🏋️‍♂️ FUERZA")
                ed_f = st.data_editor(d_fue, num_rows="dynamic", use_container_width=True, height=400, key=f"f_{alu}_{dia}", column_config=cfg_fuerza, hide_index=True)
                
                st.caption("🏃‍♂️ CARDIO")
                ed_ca = st.data_editor(d_car, num_rows="dynamic", use_container_width=True, key=f"ca_{alu}_{dia}", column_config=cfg_comun, hide_index=True)
            
            c_g_bot, c_d_bot = st.columns([1, 1])
            with c_g_bot:
                guardado_abajo = st.button("💾 GUARDAR RUTINA", type="primary", use_container_width=True, key="btn_guardar_bot")
            with c_d_bot:
                if not rut.empty: 
                    st.download_button("📥 DESCARGAR WORD", generar_word(alu, rut), f"{alu}.docx", use_container_width=True, key="btn_word_bot")

            if guardado_arriba or guardado_abajo:
                guardar_rutina_actualizada(alu, dia, ed_c, ed_f, ed_ca)
                leer_rutina.clear() 
                st.success("✅ Guardado Exitoso.")
                time.sleep(1)
                st.rerun()

        with tab2:
            us = obtener_todos_usuarios()
            als = us[us["Rol"] == "alumno"]["Usuario"].tolist()
            alu_s = st.selectbox("VER DATOS DE:", als, key="stats_sel")
            
            c_cal, c_hist = st.columns([1, 2])
            with c_cal:
                st.markdown("### 📅 Asistencia")
                dfs = leer_sesiones_alumno(alu_s)
                
                MESES_LIST = ["", "Enero", "Febrero", "Marzo", "Abril", "Mayo", "Junio", "Julio", "Agosto", "Septiembre", "Octubre", "Noviembre", "Diciembre"]
                c_sel_m, c_sel_y = st.columns([2, 1])
                mes_actual_idx = datetime.now().month
                sel_mes_nom = c_sel_m.selectbox("Mes", MESES_LIST[1:], index=mes_actual_idx-1, key="admin_month")
                sel_anio = c_sel_y.number_input("Año", value=datetime.now().year, key="admin_year")
                sel_mes_idx = MESES_LIST.index(sel_mes_nom)

                count_month = 0; count_year = 0
                if not dfs.empty:
                     count_month = dfs[(dfs["Fecha"].dt.year == sel_anio) & (dfs["Fecha"].dt.month == sel_mes_idx)]["Fecha"].nunique()
                     count_year = dfs[dfs["Fecha"].dt.year == sel_anio]["Fecha"].nunique()
                
                km1, km2 = st.columns(2)
                km1.metric(f"Mes", count_month)
                km2.metric(f"Año", count_year)

                render_calendar(int(sel_anio), sel_mes_idx, dfs)

            with c_hist:
                st.markdown("### 📈 Historial")
                df_r = leer_registros_alumno(alu_s)
                with st.expander("Gestión de Registros"):
                     if not df_r.empty:
                        cols_show = ["Fecha", "Usuario", "Ejercicio", "Peso", "Repeticiones", "Rpe", "Notas"]
                        for c in cols_show: 
                            if c not in df_r.columns: df_r[c] = ""
                        edited_hist = st.data_editor(df_r[cols_show], num_rows="dynamic", key=f"hist_edit_{alu_s}")
                        if st.button("Actualizar Historial"):
                            actualizar_registros_masivo(alu_s, edited_hist)
                            leer_registros_alumno.clear() 
                            st.rerun()
                     else: st.info("Sin registros.")
                if not df_r.empty and "Peso_Grafico" in df_r.columns:
                    lista_ejercicios = df_r["Ejercicio"].unique()
                    if len(lista_ejercicios) > 0:
                        ej_v = st.selectbox("Gráfico de:", lista_ejercicios)
                        df_plt = df_r[df_r["Ejercicio"] == ej_v].sort_values("Fecha", ascending=True)
                        if not df_plt.empty:
                            df_plt['1RM'] = df_plt['Peso_Grafico'] * (1 + (df_plt['Reps_Grafico'] / 30))
                            st.bar_chart(df_plt.set_index("Fecha")["1RM"], color="#E63946")
                            
        with tab3:
            st.markdown(f"### 👤 Perfil: {nombre.upper()}")
            st.caption("Administrador Principal")
            st.divider()
            st.info("Desde aquí puedes gestionar tu cuenta y el rendimiento de la aplicación.")
            
            if st.button("🔴 LIMPIAR CACHÉ DE LA APP", use_container_width=True): 
                st.cache_data.clear()
                st.success("Memoria borrada. Actualizando...")
                time.sleep(1)
                st.rerun()
                
            st.markdown("<br>", unsafe_allow_html=True)
            if st.button("🚪 CERRAR SESIÓN", use_container_width=True, key="btn_logout_admin"):
                cerrar_sesion_seguro()

    else:
        st.title(f"RUTINA DE {nombre.upper()}")
        
        t1, t2, t3, t4 = st.tabs(["💪 ENTRENAR", "📅 ASISTENCIA", "📝 BITÁCORA", "⚙️ PERFIL"])
        
        with t1:
            rut = leer_rutina(alias)
            if not rut.empty:
                with st.container(border=True):
                    col_d, col_w = st.columns([3, 1])
                    with col_d:
                        dias_disponibles = sorted(rut["Dia"].unique(), key=natural_sort_key)
                        d_hoy = st.selectbox("Selecciona Día", dias_disponibles, key="selector_dia_alumno")

                    with col_w:
                         st.markdown("<br>", unsafe_allow_html=True)
                         st.download_button("📥 Word", generar_word(alias, rut), "Rutina.docx", use_container_width=True)
                
                r_hoy = rut[rut["Dia"] == d_hoy]
                
                notas_importantes = []
                for idx, row in r_hoy.iterrows():
                    if str(row.get('Notas','')).strip() != "":
                        notas_importantes.append(f"**{row['Ejercicio']}**: {row['Notas']}")
                
                if notas_importantes:
                    st.info("📝 **NOTAS:**\n\n" + "\n".join([f"- {n}" for n in notas_importantes]))

                df_sec_c = r_hoy[r_hoy["Seccion"] == "Calentamiento"]
                df_sec_f = r_hoy[r_hoy["Seccion"] == "Fuerza"]
                df_sec_ca = r_hoy[r_hoy["Seccion"] == "Cardio"]

                renderizar_bloque_seccion("🔥 Calentamiento", df_sec_c, alias, d_hoy)
                renderizar_bloque_seccion("🏋️‍♂️ Fuerza", df_sec_f, alias, d_hoy)
                renderizar_bloque_seccion("🏃‍♂️ Cardio", df_sec_ca, alias, d_hoy)

                st.markdown("<br>", unsafe_allow_html=True)
                
                if st.button("💾 GUARDAR CAMBIOS EN LA RUTINA", type="primary", use_container_width=True):
                    guardar_desde_tarjetas(alias, d_hoy, r_hoy, st.session_state)
                    leer_rutina.clear()
                    st.toast("✅ Notas y Kilos guardados!")
                    time.sleep(1)
                    st.rerun()

                st.markdown("---")
                
                with st.expander("📝 Registro Rápido (Guardar serie en historial)", expanded=False):
                    with st.form("registro_rapido"):
                        st.markdown("⭐ **TODOS LOS CAMPOS SON OBLIGATORIOS**")
                        ejercicios_fuerza = r_hoy[r_hoy["Seccion"] == "Fuerza"]["Ejercicio"].unique()
                        if len(ejercicios_fuerza) == 0:
                            ej_sel = st.text_input("Ejercicio")
                        else:
                            c_ej, c_kg = st.columns([2, 1])
                            ej_sel = c_ej.selectbox("Ejercicio", ejercicios_fuerza)
                            kg_in = c_kg.text_input("Kilos", placeholder="Ej: 72.5 o 72,5")
                        
                        c_reps, c_rpe = st.columns(2)
                        reps_in = c_reps.text_input("Reps", placeholder="Ej: 10")
                        
                        st.caption("🧠 **¿Qué es el RPE?** Es tu Percepción de Esfuerzo (del 1 al 10). \n* **RPE 10:** Fallo muscular, no podías hacer ni una repetición más.\n* **RPE 8:** Esfuerzo alto, te guardaste unas 2 repeticiones en el tanque.")
                        rpe_in = c_rpe.slider("RPE", 1, 10, 7)
                        
                        notas_in = st.text_area("Notas extra", height=60)
                        
                        if st.form_submit_button("GUARDAR EN HISTORIAL", use_container_width=True):
                            if not ej_sel.strip() or not kg_in.strip() or not reps_in.strip():
                                st.error("⚠️ Error: Ejercicio, Kilos y Reps son campos obligatorios.")
                            else:
                                try:
                                    kg_limpio = float(kg_in.replace(",", "."))
                                    reps_limpio = int(reps_in.strip())
                                    
                                    guardar_registro(alias, ej_sel, kg_limpio, reps_limpio, rpe_in, notas_in)
                                    leer_registros_alumno.clear() 
                                    st.success("✅ ¡Guardado con éxito!")
                                    time.sleep(1)
                                    st.rerun()
                                except ValueError:
                                    st.error("⚠️ Error: Los Kilos y Reps deben ser solo números (Ej: 72.5). No escribas letras.")

                c_ok, c_fail = st.columns(2)
                
                # CTO FIX: Ahora cuando aprietan TERMINÉ, también guardamos en la base de datos qué "d_hoy" acaban de terminar.
                if c_ok.button("✅ TERMINÉ POR HOY", use_container_width=True):
                    guardar_estado_sesion(alias, "Completado", fecha_dt=None, dia_rutina=d_hoy)
                    leer_sesiones_alumno.clear() 
                    st.balloons()
                if c_fail.button("⚠️ INCOMPLETO", use_container_width=True):
                    guardar_estado_sesion(alias, "Incompleto", fecha_dt=None, dia_rutina=d_hoy)
                    leer_sesiones_alumno.clear()
            else: st.info("No tienes rutina asignada aún.")
            
        with t2:
            st.markdown("### 📅 Asistencia")
            dfs = leer_sesiones_alumno(alias)
            
            # Buscamos las rutinas del alumno para saber qué opciones darle (Día 1, Día 2, etc.)
            rut_asistencia = leer_rutina(alias)
            dias_opciones = sorted(rut_asistencia["Dia"].unique(), key=natural_sort_key) if not rut_asistencia.empty else ["Día 1", "Día 2", "Día 3"]
            
            # CTO FIX: La magia visual. Buscamos el último "Completado" y mostramos el cartelito.
            if not dfs.empty and "Dia_rutina" in dfs.columns:
                df_completados = dfs[dfs["Estado"] == "Completado"].sort_values("Fecha", ascending=False)
                if not df_completados.empty:
                    ultimo_registro = df_completados.iloc[0]
                    fecha_str = ultimo_registro["Fecha"].strftime("%d/%m/%Y")
                    dia_str = str(ultimo_registro.get("Dia_rutina", ""))
                    if dia_str.strip():
                        st.info(f"📍 **Último entrenamiento registrado:** {dia_str} (el {fecha_str})")

            MESES_LIST = ["", "Enero", "Febrero", "Marzo", "Abril", "Mayo", "Junio", "Julio", "Agosto", "Septiembre", "Octubre", "Noviembre", "Diciembre"]
            c_sel_m, c_sel_y = st.columns([2, 1])
            mes_actual_idx = datetime.now().month
            sel_mes_nom = c_sel_m.selectbox("Mes", MESES_LIST[1:], index=mes_actual_idx-1)
            sel_anio = c_sel_y.number_input("Año", value=datetime.now().year)
            sel_mes_idx = MESES_LIST.index(sel_mes_nom)
            
            c_kpi1, c_kpi2 = st.columns(2)
            count_month = 0; count_year = 0
            if not dfs.empty:
                 count_month = dfs[(dfs["Fecha"].dt.year == sel_anio) & (dfs["Fecha"].dt.month == sel_mes_idx)]["Fecha"].nunique()
                 count_year = dfs[dfs["Fecha"].dt.year == sel_anio]["Fecha"].nunique()
            c_kpi1.metric(f"Entrenos {sel_mes_nom}", count_month)
            c_kpi2.metric(f"Total {sel_anio}", count_year)
            render_calendar(int(sel_anio), sel_mes_idx, dfs)
            
            with st.expander("🗓️ ¿Olvidaste dar el presente?"):
                st.caption("Registra un entrenamiento que hiciste antes y olvidaste marcar.")
                
                # CTO FIX: Dividí el espacio en 3 columnas para que entre el nuevo selector de día.
                col_d_pas, col_dia_pas, col_b_pas = st.columns([2, 2, 2])
                fecha_pasada = col_d_pas.date_input("Fecha", value=datetime.now().date() - timedelta(days=1), max_value=datetime.now().date())
                
                # Agregamos la lista desplegable
                dia_pasado = col_dia_pas.selectbox("¿Qué día hiciste?", dias_opciones)
                
                if col_b_pas.button("Marcar Completado", use_container_width=True):
                    # Guardamos enviando la fecha pasada Y el día seleccionado
                    guardar_estado_sesion(alias, "Completado", fecha_dt=fecha_pasada, dia_rutina=dia_pasado)
                    st.success(f"Entreno ({dia_pasado}) del {fecha_pasada.strftime('%d/%m')} guardado.")
                    leer_sesiones_alumno.clear() 
                    time.sleep(1); st.rerun()

        with t3:
            st.markdown("### 📝 Bitácora")
            
            df_r = leer_registros_alumno(alias)
            
            if not df_r.empty and "Peso_Grafico" in df_r.columns:
                lista_ej = df_r["Ejercicio"].unique()
                if len(lista_ej) > 0:
                    ej_sel_bitacora = st.selectbox("📊 Selecciona un Ejercicio para ver tu progreso:", lista_ej, key="sel_bitacora")
                    
                    df_plt = df_r[df_r["Ejercicio"] == ej_sel_bitacora].sort_values("Fecha", ascending=True).copy()
                    
                    if not df_plt.empty:
                        df_chart = df_plt.copy().reset_index(drop=True)
                        df_chart["Fecha_Str"] = df_chart["Fecha"].dt.strftime("%d/%m")
                        
                        df_chart['Set_Num'] = df_chart.groupby('Fecha_Str').cumcount() + 1
                        df_chart['Eje_X'] = df_chart['Fecha_Str'] + " (S" + df_chart['Set_Num'].astype(str) + ")"
                        df_chart['Reps_Cartel'] = df_chart['Reps_Grafico'].astype(int).astype(str) + "r"
                        
                        base = alt.Chart(df_chart).encode(
                            x=alt.X('Eje_X:N', sort=None, title="Fecha y Serie", axis=alt.Axis(labelAngle=-45))
                        )
                        
                        barras = base.mark_bar(color="#E63946", opacity=0.9, cornerRadiusTopLeft=3, cornerRadiusTopRight=3).encode(
                            y=alt.Y('Peso_Grafico:Q', title="Kilos Levantados"),
                            tooltip=[
                                alt.Tooltip('Fecha_Str:N', title="Fecha"), 
                                alt.Tooltip('Peso_Grafico:Q', title="Kilos"), 
                                alt.Tooltip('Reps_Grafico:Q', title="Reps"), 
                                alt.Tooltip('Rpe:Q', title="RPE")
                            ]
                        )
                        
                        textos = base.mark_text(
                            dy=-10, 
                            color="white",
                            fontWeight="bold"
                        ).encode(
                            y=alt.Y('Peso_Grafico:Q'),
                            text=alt.Text('Reps_Cartel:N')
                        )
                        
                        chart = (barras + textos).properties(height=320) 
                        st.altair_chart(chart, use_container_width=True)
                        
                    else:
                        st.info("Sin datos para este ejercicio.")
                else:
                    st.info("No hay ejercicios registrados.")
                
                st.markdown("---")
                
                st.markdown(f"#### 🗑️ Gestor de Registros: {ej_sel_bitacora}")
                st.caption("Selecciona las series que anotaste mal o que están repetidas y bórralas.")
                
                df_hist_editor = df_plt.sort_values("Fecha", ascending=False)
                
                registros_a_borrar = []
                for idx, row in df_hist_editor.iterrows():
                    fecha_str = row['Fecha'].strftime('%d/%m/%Y') if pd.notnull(row['Fecha']) else "Sin Fecha"
                    peso_val = row.get('Peso', '-')
                    reps_val = row.get('Repeticiones', '-')
                    rpe_val = row.get('Rpe', '-')
                    
                    etiqueta = f"📅 {fecha_str} | ⚖️ {peso_val} kg x {reps_val} reps | 🔥 RPE: {rpe_val}"
                    if st.checkbox(etiqueta, key=f"del_{idx}"):
                        registros_a_borrar.append(idx)
                
                if st.button("🚨 BORRAR SELECCIONADOS", type="primary", use_container_width=True):
                    if len(registros_a_borrar) > 0:
                        df_plt_limpio = df_plt.drop(index=registros_a_borrar)
                        
                        df_otros_ejercicios = df_r[df_r["Ejercicio"] != ej_sel_bitacora]
                        df_a_guardar = pd.concat([df_otros_ejercicios, df_plt_limpio], ignore_index=True)
                        
                        cols_oficiales = ["Fecha", "Usuario", "Ejercicio", "Peso", "Repeticiones", "Rpe", "Notas"]
                        for c in cols_oficiales:
                            if c not in df_a_guardar.columns: df_a_guardar[c] = ""
                        
                        df_a_guardar = df_a_guardar[cols_oficiales]
                        df_a_guardar["Usuario"] = alias
                        
                        actualizar_registros_masivo(alias, df_a_guardar)
                        leer_registros_alumno.clear()
                        st.success("¡Registros eliminados de tu historial!")
                        time.sleep(1)
                        st.rerun()
                    else:
                        st.warning("No seleccionaste ninguna serie para borrar.")
            else:
                st.info("Aún no tienes registros en tu bitácora.")

            st.markdown("---")

            st.markdown("### 🧮 Calculadora de 1RM")
            with st.container(border=True):
                c_peso, c_reps = st.columns(2)
                p_in = c_peso.number_input("Peso levantado (kg)", min_value=0.0, step=1.0, value=0.0)
                r_in = c_reps.number_input("Repeticiones logradas", min_value=1, max_value=30, step=1, value=5)
            
            rm_calculo = 0
            if p_in > 0:
                rm_calculo = calcular_1rm_promedio(p_in, r_in)
            
            if rm_calculo > 0:
                with st.expander("📊 Ver Tablas de RM y % de Fuerza", expanded=True):
                    st.caption(f"TABLA DE FUERZA (Base: {int(rm_calculo)}kg)")
                    
                    md_rm = "| Repeticiones | Peso Estimado (kg) | Repeticiones | Peso Estimado (kg) |\n"
                    md_rm += "| :---: | :---: | :---: | :---: |\n"
                    for i in range(1, 7):
                        peso_izq = rm_calculo if i == 1 else rm_calculo * (1.0278 - 0.0278 * i)
                        i_der = i + 6
                        peso_der = rm_calculo * (1.0278 - 0.0278 * i_der)
                        md_rm += f"| **{i} RM** | {peso_izq:.1f} | **{i_der} RM** | {peso_der:.1f} |\n"
                    
                    st.markdown(md_rm)

                    st.markdown("#### % Cargas de Trabajo")
                    
                    with st.container(border=True):
                        st.caption("🎯 **Calculadora de % Exacto**")
                        col_calc1, col_calc2 = st.columns([1, 2])
                        porc_custom = col_calc1.number_input("Ingresa %", min_value=1.0, max_value=200.0, value=72.5, step=0.5)
                        val_custom = rm_calculo * (porc_custom / 100)
                        col_calc2.info(f"🔥 **{porc_custom}%** = {val_custom:.1f} kg")
                    
                    st.markdown("<br>", unsafe_allow_html=True)
                    
                    porcentajes_izq = [125, 115, 105, 95, 85, 75, 65, 55, 45, 35]
                    porcentajes_der = [120, 110, 100, 90, 80, 70, 60, 50, 40, 30]
                    
                    md_porc = "| % Fuerza | Carga (kg) | % Fuerza | Carga (kg) |\n"
                    md_porc += "| :---: | :---: | :---: | :---: |\n"
                    for p_i, p_d in zip(porcentajes_izq, porcentajes_der):
                        val_i = rm_calculo * (p_i / 100)
                        val_d = rm_calculo * (p_d / 100)
                        md_porc += f"| **{p_i}%** | {val_i:.1f} | **{p_d}%** | {val_d:.1f} |\n"
                        
                    st.markdown(md_porc)
                        
        with t4:
            st.markdown(f"### 👤 Perfil: {nombre}")
            st.caption(f"Usuario activo: {alias}")
            st.divider()
            
            st.info("💡 Desde aquí puedes cerrar tu sesión de forma segura si compartes este dispositivo o deseas ingresar con otra cuenta.")
            
            if st.button("🚪 CERRAR SESIÓN", use_container_width=True, key="btn_logout_alu"):
                cerrar_sesion_seguro()
